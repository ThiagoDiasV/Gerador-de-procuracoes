# gerador_procuracao.py
# Criador de procurações em Python
import xlsxwriter
from docx import Document
import xlrd
from faker import Faker
from random import randint
from os import listdir


def create_worksheet(person_data):
    """
    This function creates a worksheet and put the data on it
    """
    # Cria o arquivo .xlsx
    workbook = xlsxwriter.Workbook('vizviz.xlsx')

    # Cria a planilha dentro do arquivo
    worksheet = workbook.add_worksheet()

    # Cria os nomes das colunas
    labels = (
        'Nome Outorgante', 'Nacionalidade Outorgante',
        'Estado Civil Outorgante', 'Profissão Outorgante',
        'CPF Outorgante', 'RG Outorgante', 'Endereço Outorgante',
        'Nome Outorgado', 'Nacionalidade Outorgado', 'Estado Civil Outorgado',
        'Profissão Outorgado', 'CPF Outorgado', 'RG Outorgado',
        'Endereço Outorgado', 'Objetivo procuração', 'Validade Procuração'
    )

    # Escreve os nomes das colunas na planilha
    worksheet.write_row('A1', labels)

    # Coloca os dados das pessoas na planilha
    for i, person in enumerate(person_data.values()):
        worksheet.write_row(f'A{i+2}', person)

    # Fecha o arquivo
    workbook.close()


def create_data():
    """
    This function creates the data of each person using an instance of Faker
    """

    # Cria uma instância de Faker, com dados do Brasil
    faker = Faker('pt_BR')

    # Cria o dicionário no qual serão adicionados os dados
    data = dict()

    # Cria duas listas com dados de estado civil
    # A lib Faker não dá suporte para esse dado
    marital_status_male = ['Solteiro', 'Casado', 'Divorciado', 'Viúvo']
    marital_status_female = ['Solteira', 'Casada', 'Divorciada', 'Viúva']

    # Cria os dados usando duas possibilidades
    # Se random_number == 1, os nomes são masculinos
    # Se random_number == 0, os nomes são femininos
    for i in range(11):
        random_number = randint(0, 1)
        if random_number:
            data[f'Person{i}'] = [
                faker.name_male(), 'Brasileiro',
                marital_status_male[randint(0, len(marital_status_male)-1)],
                faker.job(),
                faker.pyint(min_value=10000000000, max_value=99999999999),
                faker.pyint(min_value=1000000000000, max_value=9999999999999),
                faker.address(), faker.name_male(), 'Brasileiro',
                marital_status_male[randint(0, len(marital_status_male)-1)],
                faker.job(),
                faker.pyint(min_value=10000000000, max_value=99999999999),
                faker.pyint(min_value=1000000000000, max_value=9999999999999),
                faker.address().replace('\n', ' '),
                f'Representar perante a empresa {faker.company()} para blabla',
                faker.date(pattern='%d-%m-%Y', end_datetime='+20y')
                ]
        else:
            data[f'Person{i}'] = [
                faker.name_female(), 'Brasileira',
                marital_status_female[randint(0, len(marital_status_female)-1)],
                faker.job(),
                faker.pyint(min_value=10000000000, max_value=99999999999),
                faker.pyint(min_value=1000000000000, max_value=9999999999999),
                faker.address().replace('\n', ' '), faker.name_female(), 'Brasileira',
                marital_status_female[randint(0, len(marital_status_female)-1)],
                faker.job(),
                faker.pyint(min_value=10000000000, max_value=99999999999),
                faker.pyint(min_value=1000000000000, max_value=9999999999999),
                faker.address(),
                f'representar perante a empresa {faker.company()} para blabla',
                faker.date(pattern='%d-%m-%Y', end_datetime='+20y')
                ]
    return data


def read_xlsx():
    """
    This function read the created workbook
    """
    # Pega o arquivo .xlsx no diretório atual
    excel_file = [i for i in listdir('.') if i.endswith('.xlsx')][0]

    # Abre o arquivo usando a lib xlrd
    book = xlrd.open_workbook(excel_file)

    # Pega a primeira planilha do arquivo
    sheet = book.sheet_by_index(0)

    # Adiciona os dados de cada pessoa a um dicionário
    data = dict()
    for i in range(1, sheet.nrows):
        data[f'Person{i}'] = sheet.row_values(i)

    return data


def create_docx(person_data):
    """
    This function creates a template in docx file and add the data to it
    for each person
    """
    
    # Itera sobre os dados de cada pessoa, criando um documento para cada um
    for person in person_data.values():

        # Transforma os dados de cada pessoa em uma lista
        data = list(person)

        # Cria o documento
        document = Document()

        # Adiciona um título
        document.add_heading('Procuração')
        
        # Primeiro parágrafo
        paragraph_1 = f'''Outorgante: Eu, {data[0]}, {data[1]}, {data[2]}, {data[3]}, portador (a) do CPF nº {int(data[4])}, e do RG nº {int(data[5])}, residente e domicilado (a) em {data[6]}, pelo presente instrumento, nomeio como meu (minha) procurador (a)'''

        # Segundo parágrafo
        paragraph_2 = f'''Outorgado (a): {data[7]}, {data[8]}, {data[9]}, {data[10]}, portador (a) do CPF nº {int(data[11])}, e do RG nº {int(data[12])}, residente e domiciliado (a) em {data[13]}, com poderes para representar o outorgante com objetivo de {data[14]}, responsabilizando-me por todos os atos praticados no cumprimento deste instrumento, cessando seus efeitos em {data[15]}.
        
        Local, Data





        Assinatura
        '''

        # Adiciona os parágrafos ao documento
        document.add_paragraph(paragraph_1)
        document.add_paragraph(paragraph_2)

        # Salva o documento
        document.save(f'Procuração de {data[0]}.docx')



data = create_data()
table = create_worksheet(data)
new_data = read_xlsx()
create_docx(new_data)
